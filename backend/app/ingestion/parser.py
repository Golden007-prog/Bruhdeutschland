"""Parse + normalize uploads into the typed :class:`Profile` (CLAUDE.md §2 rule 7; build Phase 4).

Three sources: a structured intake form (primary, fully typed), a LinkedIn JSON export, and a
resume file (PDF / DOCX / text). Large documents are chunked; PII is redacted before anything is
logged. No official facts are produced here — only the student's own data is normalized.
"""

from __future__ import annotations

import io
from typing import Any

from pydantic import BaseModel, Field

from app.core.logging import redact_pii
from app.llm.schemas import EducationItem, Profile
from app.models.enums import ProfileSourceType
from app.services.ects_calculator import CourseCredit
from app.services.gpa_converter import GradeScale

MAX_CHUNK_CHARS = 4000


class UnsupportedDocument(ValueError):
    """Raised when a resume file type cannot be parsed."""


class GpaScaleForm(BaseModel):
    best: float
    min_pass: float


class CourseForm(BaseModel):
    credits: float = Field(ge=0)
    name: str | None = None
    passed: bool = True


class IntakeForm(BaseModel):
    """User-facing intake payload (friendly field names). Normalized into :class:`Profile`."""

    full_name: str | None = None
    education: list[EducationItem] = []
    experience: list[str] = []
    skills: list[str] = []
    raw_gpa: float | None = None
    gpa_scale: GpaScaleForm | None = None
    courses: list[CourseForm] = []
    goals: list[str] = []
    field: str | None = None
    degree_level: str | None = None
    target_intake: str | None = None


def parse_intake(form: IntakeForm) -> Profile:
    """Normalize a structured intake form into a typed Profile."""
    return Profile(
        source_type=ProfileSourceType.INTAKE,
        full_name=form.full_name,
        education=form.education,
        experience=form.experience,
        skills=form.skills,
        raw_gpa=form.raw_gpa,
        gpa_scale=(
            GradeScale(best=form.gpa_scale.best, min_pass=form.gpa_scale.min_pass)
            if form.gpa_scale
            else None
        ),
        courses=[
            CourseCredit(credits=c.credits, name=c.name, passed=c.passed) for c in form.courses
        ],
        goals=form.goals,
        field=form.field,
        degree_level=form.degree_level,
        target_intake=form.target_intake,
    )


def parse_linkedin(export: dict[str, Any]) -> Profile:
    """Map a LinkedIn export (subset of fields) into a typed Profile."""
    edu = [
        EducationItem(
            institution=e.get("school", e.get("institution", "Unknown")),
            degree=e.get("degree"),
            field=e.get("field"),
            year=e.get("end_year"),
        )
        for e in export.get("education", [])
    ]
    positions = [
        f"{p.get('title', '?')} @ {p.get('company', '?')}" for p in export.get("positions", [])
    ]
    return Profile(
        source_type=ProfileSourceType.LINKEDIN,
        full_name=export.get("full_name"),
        education=edu,
        experience=positions,
        skills=list(export.get("skills", [])),
        field=export.get("field"),
    )


def chunk_text(text: str, *, max_chars: int = MAX_CHUNK_CHARS) -> list[str]:
    """Split text into <= max_chars chunks on paragraph boundaries (for embedding/LLM extraction)."""
    if not text.strip():
        return []
    chunks: list[str] = []
    current = ""
    for para in text.split("\n\n"):
        if len(current) + len(para) + 2 > max_chars and current:
            chunks.append(current.strip())
            current = ""
        current += para + "\n\n"
    if current.strip():
        chunks.append(current.strip())
    return chunks


def extract_text(data: bytes, *, content_type: str, filename: str) -> str:
    """Extract plain text from a resume upload (PDF, DOCX, or text)."""
    name = filename.lower()
    if content_type == "application/pdf" or name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx") or "officedocument.wordprocessing" in content_type:
        import docx

        document = docx.Document(io.BytesIO(data))
        return "\n\n".join(p.text for p in document.paragraphs)
    if content_type.startswith("text/") or name.endswith((".txt", ".md")):
        return data.decode("utf-8", errors="replace")
    raise UnsupportedDocument(
        f"Cannot parse '{filename}' ({content_type}). Upload PDF, DOCX, text, or use the intake form."
    )


def parse_resume(data: bytes, *, content_type: str, filename: str) -> tuple[Profile, list[str]]:
    """Extract + chunk a resume into a skeleton Profile and its text chunks.

    Deep structured extraction (education/skills/grades) is an LLM task layered on top; this step is
    deterministic. The full name is taken from the first non-empty line; raw text is never logged.
    """
    text = extract_text(data, content_type=content_type, filename=filename)
    chunks = chunk_text(text)
    first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), None)
    full_name = redact_pii(first_line) if first_line and len(first_line) <= 80 else None
    profile = Profile(source_type=ProfileSourceType.RESUME, full_name=full_name)
    return profile, chunks
